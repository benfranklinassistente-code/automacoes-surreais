#!/usr/bin/env python3
"""
Converte ebook para DOCX formatado profissionalmente
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def criar_ebook_docx():
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    
    # === CAPA ===
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('\n\n\n\n')
    run = titulo.add_run('📱 GUIA DE EMERGÊNCIA')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run('CONTA HACKEADA')
    run.bold = True
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor(118, 75, 162)
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('\nRecupere seu WhatsApp em 5 Passos Simples')
    run.font.size = Pt(18)
    run.italic = True
    
    doc.add_paragraph('\n\n\n')
    
    protecao = doc.add_paragraph()
    protecao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = protecao.add_run('🛡️ Proteção para Idosos 60+')
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph('\n\n\n\n')
    
    by = doc.add_paragraph()
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = by.add_run('Por')
    run.font.size = Pt(14)
    
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo.add_run('60maisPlay')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(255, 193, 7)
    
    data = doc.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = data.add_run('Fevereiro 2026')
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # === SUMÁRIO ===
    heading = doc.add_heading('SUMÁRIO', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    sumario_itens = [
        '1. Introdução',
        '2. O Problema',
        '3. Solução Passo a Passo',
        '4. Checklist de Prevenção',
        '5. Recursos Adicionais'
    ]
    
    for item in sumario_itens:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # === INTRODUÇÃO ===
    heading = doc.add_heading('INTRODUÇÃO', level=1)
    heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph(
        'Olá! Se você está lendo este guia, provavelmente já passou por ter sua conta '
        'clonada ou hackeada, ou conhece alguém que passou por isso.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('A boa notícia é que existe solução e ela é mais simples do que parece!')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(40, 167, 69)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('Neste guia prático, você vai aprender:')
    
    aprende = [
        'Como identificar o problema rapidamente',
        'Os 5 passos para resolver',
        'Como prevenir que aconteça de novo'
    ]
    for item in aprende:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    doc.add_paragraph()
    
    # Box de tempo
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Inches(0.5)
    box.paragraph_format.right_indent = Inches(0.5)
    run = box.add_run('⏱️ Tempo estimado de leitura: 10 minutos\n')
    run.italic = True
    run = box.add_run('⚡ Tempo para aplicar: 15 minutos')
    run.italic = True
    
    doc.add_page_break()
    
    # === O PROBLEMA ===
    heading = doc.add_heading('O PROBLEMA', level=1)
    heading.runs[0].font.color.rgb = RGBColor(220, 53, 69)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'O golpe da clonagem de WhatsApp afeta milhares de idosos todos os dias no Brasil. '
        'Os criminosos usam engenharia social para conseguir o código de verificação e tomar '
        'controle da conta.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('SINAIS DE ALERTA')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(255, 193, 7)
    
    doc.add_paragraph('Fique atento a estes sinais:')
    
    sinais = [
        'Você recebeu uma mensagem de um amigo pedindo o código que você "recebeu por engano"',
        'Seus contatos dizem que estão recebendo mensagens suspeitas de você',
        'Foi desconectado do WhatsApp de repente',
        'Alguém ligou dizendo ser do "suporte do WhatsApp"'
    ]
    for sinal in sinais:
        p = doc.add_paragraph(f'□ {sinal}', style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_page_break()
    
    # === SOLUÇÃO ===
    heading = doc.add_heading('SOLUÇÃO PASSO A PASSO', level=1)
    heading.runs[0].font.color.rgb = RGBColor(40, 167, 69)
    
    doc.add_paragraph()
    
    # PASSO 1
    p = doc.add_paragraph()
    run = p.add_run('PASSO 1: NÃO ENTRE EM PÂNICO')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    doc.add_paragraph('O que fazer:').runs[0].bold = True
    doc.add_paragraph('• Respire fundo', style='List Bullet')
    doc.add_paragraph('• Lembre-se: existe solução', style='List Bullet')
    doc.add_paragraph('• Vamos resolver juntos passo a passo', style='List Bullet')
    
    doc.add_paragraph()
    
    # Dica box
    dica = doc.add_paragraph()
    dica.paragraph_format.left_indent = Inches(0.5)
    run = dica.add_run('💡 Dica do Vovô:\n')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 193, 7)
    run = dica.add_run('"Calma é a melhor arma contra o desespero. Um passo de cada vez!"')
    run.italic = True
    
    doc.add_paragraph()
    
    # PASSO 2
    p = doc.add_paragraph()
    run = p.add_run('PASSO 2: RECUPERE O ACESSO IMEDIATAMENTE')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    doc.add_paragraph('O que fazer:').runs[0].bold = True
    doc.add_paragraph('• Abra o WhatsApp no seu celular', style='List Bullet')
    doc.add_paragraph('• Toque em "Verificar" quando pedir o código', style='List Bullet')
    doc.add_paragraph('• Peça para reenviar o código por SMS', style='List Bullet')
    
    p = doc.add_paragraph('• NUNCA compartilhe este código com ninguém', style='List Bullet')
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(220, 53, 69)
    
    doc.add_paragraph()
    
    # PASSO 3
    p = doc.add_paragraph()
    run = p.add_run('PASSO 3: AVISE SEUS CONTATOS')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    doc.add_paragraph('O que fazer:').runs[0].bold = True
    doc.add_paragraph('• Envie uma mensagem para seus grupos e contatos importantes', style='List Bullet')
    doc.add_paragraph('• Explique que sua conta foi clonada', style='List Bullet')
    doc.add_paragraph('• Peça para ignorarem mensagens suspeitas', style='List Bullet')
    doc.add_paragraph('• Peça para não clicarem em links', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Mensagem sugerida:')
    run.bold = True
    run.italic = True
    
    msg = doc.add_paragraph()
    msg.paragraph_format.left_indent = Inches(0.5)
    msg.paragraph_format.right_indent = Inches(0.5)
    run = msg.add_run('"Olá! Minha conta do WhatsApp foi clonada. Por favor, ignore qualquer '
                     'mensagem estranha que você tenha recebido de mim. Já recuperei minha conta. Obrigado!"')
    run.italic = True
    run.font.color.rgb = RGBColor(108, 117, 125)
    
    doc.add_paragraph()
    
    # PASSO 4
    p = doc.add_paragraph()
    run = p.add_run('PASSO 4: ATIVE A VERIFICAÇÃO EM DUAS ETAPAS')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    doc.add_paragraph('O que fazer:').runs[0].bold = True
    doc.add_paragraph('• Vá em Configurações → Conta → Verificação em duas etapas', style='List Bullet')
    doc.add_paragraph('• Ative a opção', style='List Bullet')
    doc.add_paragraph('• Crie um PIN de 6 dígitos (guarde em lugar seguro!)', style='List Bullet')
    doc.add_paragraph('• Adicione um email de recuperação', style='List Bullet')
    
    doc.add_paragraph()
    
    # PASSO 5
    p = doc.add_paragraph()
    run = p.add_run('PASSO 5: REFORCE A SEGURANÇA')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    doc.add_paragraph('O que fazer:').runs[0].bold = True
    doc.add_paragraph('• Verifique se seus contatos estão salvos', style='List Bullet')
    doc.add_paragraph('• Faça backup das conversas importantes', style='List Bullet')
    doc.add_paragraph('• Configure o WhatsApp Web apenas quando precisar', style='List Bullet')
    doc.add_paragraph('• Desconecte todos os dispositivos vinculados', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Verificação:')
    run.bold = True
    
    checks = [
        'Recuperei o acesso à conta',
        'Avisi meus contatos',
        'Ativei verificação em duas etapas',
        'Reforcei minha segurança'
    ]
    for check in checks:
        p = doc.add_paragraph(f'✅ {check}', style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(40, 167, 69)
    
    doc.add_page_break()
    
    # === CHECKLIST ===
    heading = doc.add_heading('CHECKLIST DE PREVENÇÃO', level=1)
    heading.runs[0].font.color.rgb = RGBColor(255, 193, 7)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph(
        'Imprima esta página e deixe na geladeira ou perto do computador:'
    )
    p.runs[0].italic = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('ANTES DE QUALQUER TRANSAÇÃO:')
    run.bold = True
    run.font.size = Pt(14)
    
    prevencao = [
        'Desconfie de mensagens pedindo dinheiro emergencial',
        'Confirme por ligação de voz ou vídeo antes de transferir',
        'Nunca clique em links de remetentes desconhecidos',
        'Desconfie de ofertas muito boas para ser verdade'
    ]
    for item in prevencao:
        doc.add_paragraph(f'□ {item}', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('CONFIGURAÇÕES DE SEGURANÇA:')
    run.bold = True
    run.font.size = Pt(14)
    
    config = [
        'Verificação em duas etapas: ATIVADA',
        'Email de recuperação cadastrado',
        'Backup automático configurado',
        'PIN de segurança memorizado'
    ]
    for item in config:
        doc.add_paragraph(f'□ {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # === RECURSOS ===
    heading = doc.add_heading('RECURSOS ADICIONAIS', level=1)
    heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    # Vídeo
    p = doc.add_paragraph()
    run = p.add_run('🎥 VÍDEO TUTORIAL')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph('Assista o passo a passo em vídeo completo:')
    p = doc.add_paragraph('Link: https://60maisplay.com.br/tutorial-seguranca')
    p.runs[0].font.color.rgb = RGBColor(0, 123, 255)
    
    doc.add_paragraph()
    
    # Suporte
    p = doc.add_paragraph()
    run = p.add_run('📞 SUPORTE')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph('Dúvidas? Fale conosco:')
    doc.add_paragraph('• WhatsApp: (11) 95354-5939', style='List Bullet')
    doc.add_paragraph('• Email: benjamin@60maiscursos.com.br', style='List Bullet')
    p = doc.add_paragraph('• Site: https://60maisplay.com.br', style='List Bullet')
    p.runs[0].font.color.rgb = RGBColor(0, 123, 255)
    
    doc.add_paragraph()
    
    # Curso
    p = doc.add_paragraph()
    run = p.add_run('🎓 CURSO COMPLETO')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(118, 75, 162)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('Quer dominar a tecnologia com segurança?')
    p.runs[0].italic = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Conheça nosso curso completo:\n')
    run.font.size = Pt(14)
    run = p.add_run('"ESCUDO ANTI-GOLPES 60+"')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(220, 53, 69)
    
    doc.add_paragraph()
    
    doc.add_paragraph('Conteúdo:').runs[0].bold = True
    conteudo = [
        '10 módulos de segurança digital',
        'Aulas em vídeo simples',
        'Material de apoio impresso',
        'Grupo VIP de alunos',
        'Suporte por 6 meses'
    ]
    for item in conteudo:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Investimento: R$ 47,00')
    run.bold = True
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    run = p.add_run('Garantia: 7 dias de garantia incondicional')
    run.italic = True
    run.font.color.rgb = RGBColor(40, 167, 69)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('Para se inscrever, acesse:')
    p.runs[0].bold = True
    p = doc.add_paragraph('https://60maisplay.com.br/escudo-anti-golpes')
    p.runs[0].font.color.rgb = RGBColor(0, 123, 255)
    
    doc.add_page_break()
    
    # === SOBRE ===
    heading = doc.add_heading('SOBRE O 60maisPlay', level=1)
    heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Somos a plataforma de tecnologia para idosos 60+. Nossa missão é tornar a '
        'tecnologia acessível, segura e descomplicada para você.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Já ajudamos mais de 500 idosos')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(40, 167, 69)
    p.add_run(' a usarem tecnologia com confiança e segurança.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Nossos valores:').runs[0].bold = True
    valores = [
        'Linguagem simples, sem termos técnicos',
        'Respeito ao seu ritmo de aprendizado',
        'Suporte humano e paciente',
        'Conteúdo atualizado constantemente'
    ]
    for valor in valores:
        doc.add_paragraph(f'⭐ {valor}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('© 2026 60maisPlay. Todos os direitos reservados.\n\n')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(108, 117, 125)
    
    run = footer.add_run('Este material é gratuito. Sinta-se à vontade para compartilhar '
                        'com amigos e familiares.')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(108, 117, 125)
    
    # Salvar
    doc.save('/root/.openclaw/workspace/automacoes-surreais/4-lead-magnet/output/Guia-Conta-Hackeada-60mais.docx')
    print('✅ Ebook DOCX criado com sucesso!')
    print('📁 Local: /root/.openclaw/workspace/automacoes-surreais/4-lead-magnet/output/Guia-Conta-Hackeada-60mais.docx')

if __name__ == '__main__':
    criar_ebook_docx()
